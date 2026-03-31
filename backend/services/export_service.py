"""
导出服务
支持将小说摘要和标签导出为 TXT 或 DOCX 格式，支持批量导出
"""

import os
import zipfile
from datetime import datetime
from typing import Optional, List

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from backend.models.novel import Novel
from backend.models.summary import Summary
from backend.models.tag import Tag
from backend.models.novel_tag import NovelTag
from backend.core.config import settings


# 导出文件临时目录
EXPORT_DIR = os.path.join(settings.UPLOAD_DIR, "exports")


def _ensure_export_dir():
    """确保导出目录存在"""
    os.makedirs(EXPORT_DIR, exist_ok=True)


async def _get_novel_export_data(db: AsyncSession, novel_id: int) -> Optional[dict]:
    """
    获取小说的导出数据，包括基本信息、摘要、标签

    Returns:
        包含 title、author、summaries、tags 的字典
    """
    # 查询小说
    stmt = select(Novel).where(Novel.id == novel_id)
    result = await db.execute(stmt)
    novel = result.scalar_one_or_none()
    if not novel:
        return None

    # 查询摘要（排除分块中间摘要）
    summary_stmt = (
        select(Summary)
        .where(Summary.novel_id == novel_id, Summary.is_chunk_summary == False)
        .order_by(Summary.created_at.desc())
    )
    summary_result = await db.execute(summary_stmt)
    summaries = summary_result.scalars().all()

    # 查询标签
    tag_stmt = (
        select(Tag, NovelTag)
        .join(NovelTag, NovelTag.tag_id == Tag.id)
        .where(NovelTag.novel_id == novel_id)
        .order_by(Tag.dimension, Tag.name)
    )
    tag_result = await db.execute(tag_stmt)
    tag_rows = tag_result.all()

    # 按维度分组标签
    tags_by_dimension: dict = {}
    dimension_names = {
        "genre": "题材",
        "style": "风格",
        "element": "核心元素",
        "character": "人物类型",
        "exclusive": "专属标签",
    }
    for tag, novel_tag in tag_rows:
        dim_label = dimension_names.get(tag.dimension, tag.dimension)
        if dim_label not in tags_by_dimension:
            tags_by_dimension[dim_label] = []
        tags_by_dimension[dim_label].append(tag.name)

    return {
        "title": novel.title,
        "author": novel.author or "未知",
        "summaries": [s.content for s in summaries],
        "tags_by_dimension": tags_by_dimension,
    }


def _generate_safe_filename(title: str) -> str:
    """
    生成安全的文件名，移除不合法字符
    """
    # 移除文件名中的非法字符
    safe_title = "".join(c for c in title if c not in r'\/:*?"<>|')
    safe_title = safe_title.strip()[:50]  # 限制长度
    if not safe_title:
        safe_title = "小说"
    return safe_title


async def export_novel_txt(db: AsyncSession, novel_id: int) -> Optional[tuple]:
    """
    导出小说摘要和标签为 TXT 格式

    Returns:
        (文件路径, 文件名) 或 None
    """
    _ensure_export_dir()

    data = await _get_novel_export_data(db, novel_id)
    if not data:
        return None

    # 构建文本内容
    lines = []
    lines.append(f"小说名称：{data['title']}")
    lines.append(f"作者：{data['author']}")
    lines.append("=" * 50)
    lines.append("")

    # 摘要部分
    lines.append("【摘要】")
    if data["summaries"]:
        for i, summary in enumerate(data["summaries"], 1):
            if len(data["summaries"]) > 1:
                lines.append(f"--- 摘要 {i} ---")
            lines.append(summary)
            lines.append("")
    else:
        lines.append("暂无摘要")
        lines.append("")

    # 标签部分
    lines.append("【标签】")
    if data["tags_by_dimension"]:
        for dim, tag_names in data["tags_by_dimension"].items():
            lines.append(f"  {dim}：{', '.join(tag_names)}")
    else:
        lines.append("暂无标签")

    content = "\n".join(lines)

    # 生成文件
    date_str = datetime.now().strftime("%Y%m%d")
    safe_title = _generate_safe_filename(data["title"])
    file_name = f"{safe_title}_总结标签_{date_str}.txt"
    file_path = os.path.join(EXPORT_DIR, file_name)

    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)

    return file_path, file_name


async def export_novel_docx(db: AsyncSession, novel_id: int) -> Optional[tuple]:
    """
    导出小说摘要和标签为 DOCX 格式
    使用 python-docx 生成结构化文档

    Returns:
        (文件路径, 文件名) 或 None
    """
    _ensure_export_dir()

    data = await _get_novel_export_data(db, novel_id)
    if not data:
        return None

    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise ImportError("导出 DOCX 需要安装 python-docx: pip install python-docx")

    doc = Document()

    # 标题
    doc.add_heading(data["title"], level=1)
    doc.add_paragraph(f"作者：{data['author']}")

    # 摘要部分
    doc.add_heading("摘要", level=2)
    if data["summaries"]:
        for i, summary in enumerate(data["summaries"], 1):
            if len(data["summaries"]) > 1:
                doc.add_paragraph(f"摘要 {i}", style="Heading 3")
            doc.add_paragraph(summary)
    else:
        doc.add_paragraph("暂无摘要")

    # 标签部分（表格形式）
    doc.add_heading("标签", level=2)
    if data["tags_by_dimension"]:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        # 表头
        header_cells = table.rows[0].cells
        header_cells[0].text = "维度"
        header_cells[1].text = "标签"
        # 数据行
        for dim, tag_names in data["tags_by_dimension"].items():
            row_cells = table.add_row().cells
            row_cells[0].text = dim
            row_cells[1].text = ", ".join(tag_names)
    else:
        doc.add_paragraph("暂无标签")

    # 保存文件
    date_str = datetime.now().strftime("%Y%m%d")
    safe_title = _generate_safe_filename(data["title"])
    file_name = f"{safe_title}_总结标签_{date_str}.docx"
    file_path = os.path.join(EXPORT_DIR, file_name)

    doc.save(file_path)
    return file_path, file_name


async def batch_export(
    db: AsyncSession,
    novel_ids: List[int],
    format: str = "txt",
) -> Optional[tuple]:
    """
    批量导出多本小说的摘要和标签

    将多本小说分别导出后打包为 ZIP 文件

    Args:
        db: 数据库会话
        novel_ids: 小说 ID 列表
        format: 导出格式（txt / docx）

    Returns:
        (ZIP 文件路径, ZIP 文件名) 或 None
    """
    _ensure_export_dir()

    exported_files = []

    for novel_id in novel_ids:
        if format == "docx":
            result = await export_novel_docx(db, novel_id)
        else:
            result = await export_novel_txt(db, novel_id)

        if result:
            exported_files.append(result)

    if not exported_files:
        return None

    # 打包为 ZIP 文件
    date_str = datetime.now().strftime("%Y%m%d_%H%M%S")
    zip_name = f"批量导出_{date_str}.zip"
    zip_path = os.path.join(EXPORT_DIR, zip_name)

    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for file_path, file_name in exported_files:
            zf.write(file_path, file_name)

    # 清理单个导出文件（已打包进 ZIP）
    for file_path, _ in exported_files:
        try:
            os.remove(file_path)
        except OSError:
            pass

    return zip_path, zip_name
